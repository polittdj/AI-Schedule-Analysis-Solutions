"""Unit tests for the Phase 5 export modules.

Each test builds a realistic engine-results dict (same helper as the
Phase 3 AI tests) and verifies that the corresponding generator
produces valid bytes for the target format. "Valid" means:

* correct magic bytes (PK for zip-based docx/xlsx, %PDF for pdf)
* non-trivial size
* round-trip through the native library (python-docx Document,
  openpyxl load_workbook) without raising
* contains expected strings / sheet names so we catch regressions
  where a section silently gets dropped
"""
from __future__ import annotations

import io
from datetime import datetime
from typing import Dict, List, Optional

import pytest
from docx import Document
from openpyxl import load_workbook

from app.engine.comparator import compare_schedules
from app.engine.cpm import compute_cpm
from app.engine.dcma import compute_dcma
from app.engine.delay_analysis import analyze_delays
from app.engine.earned_value import compute_earned_value
from app.engine.float_analysis import analyze_float
from app.engine.manipulation import detect_manipulations
from app.export.docx_report import generate_docx_report
from app.export.pdf_report import generate_pdf_report
from app.export.xlsx_export import generate_xlsx_export
from app.parser.schema import (
    AssignmentData,
    ProjectInfo,
    Relationship,
    ScheduleData,
    TaskData,
)


# --------------------------------------------------------------------------- #
# Fixture
# --------------------------------------------------------------------------- #


def _make_task(
    uid: int,
    name: str,
    duration: Optional[float] = None,
    start: Optional[datetime] = None,
    finish: Optional[datetime] = None,
    baseline_start: Optional[datetime] = None,
    baseline_finish: Optional[datetime] = None,
    baseline_duration: Optional[float] = None,
    percent_complete: Optional[float] = 0.0,
    total_slack: Optional[float] = None,
    critical: bool = False,
    predecessors: Optional[List[int]] = None,
    successors: Optional[List[int]] = None,
    notes: Optional[str] = None,
    wbs: Optional[str] = None,
) -> TaskData:
    return TaskData(
        uid=uid,
        id=uid,
        name=name,
        wbs=wbs,
        duration=duration,
        start=start,
        finish=finish,
        baseline_start=baseline_start,
        baseline_finish=baseline_finish,
        baseline_duration=baseline_duration,
        percent_complete=percent_complete,
        total_slack=total_slack,
        critical=critical,
        predecessors=predecessors or [],
        successors=successors or [],
        notes=notes,
    )


def _build_comparison_results() -> Dict:
    """Build a comparative engine-results dict exercising every module."""
    prior_tasks = [
        _make_task(
            1,
            "Mobilization",
            duration=5.0,
            start=datetime(2026, 1, 5),
            finish=datetime(2026, 1, 9),
            baseline_start=datetime(2026, 1, 5),
            baseline_finish=datetime(2026, 1, 9),
            baseline_duration=5.0,
            percent_complete=100.0,
            total_slack=0.0,
            critical=True,
            successors=[2],
            wbs="1.1.1",
        ),
        _make_task(
            2,
            "Excavation",
            duration=8.0,
            start=datetime(2026, 1, 12),
            finish=datetime(2026, 1, 21),
            baseline_start=datetime(2026, 1, 12),
            baseline_finish=datetime(2026, 1, 21),
            baseline_duration=8.0,
            percent_complete=50.0,
            total_slack=0.0,
            critical=True,
            predecessors=[1],
            successors=[3],
            wbs="1.1.2",
        ),
        _make_task(
            3,
            "Concrete pour",
            duration=10.0,
            start=datetime(2026, 1, 22),
            finish=datetime(2026, 2, 4),
            baseline_start=datetime(2026, 1, 22),
            baseline_finish=datetime(2026, 2, 4),
            baseline_duration=10.0,
            percent_complete=0.0,
            total_slack=0.0,
            critical=True,
            predecessors=[2],
            successors=[4],
            notes="Weather-sensitive primary slab pour.",
            wbs="1.1.3",
        ),
        _make_task(
            4,
            "Curing",
            duration=7.0,
            start=datetime(2026, 2, 5),
            finish=datetime(2026, 2, 13),
            baseline_start=datetime(2026, 2, 5),
            baseline_finish=datetime(2026, 2, 13),
            baseline_duration=7.0,
            percent_complete=0.0,
            total_slack=0.0,
            critical=True,
            predecessors=[3],
            wbs="1.1.4",
        ),
    ]
    later_tasks = [t.model_copy(deep=True) for t in prior_tasks]
    # Slip task 3 by 5 days (weather).
    later_tasks[2] = later_tasks[2].model_copy(
        update={
            "start": datetime(2026, 1, 27),
            "finish": datetime(2026, 2, 9),
            "notes": "Delayed 5 days by storm system. Site shut down.",
        }
    )
    later_tasks[3] = later_tasks[3].model_copy(
        update={
            "start": datetime(2026, 2, 10),
            "finish": datetime(2026, 2, 18),
        }
    )

    rels = [
        Relationship(predecessor_uid=1, successor_uid=2),
        Relationship(predecessor_uid=2, successor_uid=3),
        Relationship(predecessor_uid=3, successor_uid=4),
    ]

    prior = ScheduleData(
        project_info=ProjectInfo(
            name="Bridge Replacement 42",
            status_date=datetime(2026, 1, 19),
            start_date=datetime(2026, 1, 5),
            finish_date=datetime(2026, 2, 13),
        ),
        tasks=prior_tasks,
        relationships=rels,
        assignments=[
            AssignmentData(task_uid=1, resource_uid=1, cost=1000.0, actual_cost=1000.0),
            AssignmentData(task_uid=2, resource_uid=1, cost=2000.0, actual_cost=1000.0),
        ],
    )
    later = ScheduleData(
        project_info=ProjectInfo(
            name="Bridge Replacement 42",
            status_date=datetime(2026, 1, 26),
            start_date=datetime(2026, 1, 5),
            finish_date=datetime(2026, 2, 18),
        ),
        tasks=later_tasks,
        relationships=rels,
        assignments=[
            AssignmentData(task_uid=1, resource_uid=1, cost=1000.0, actual_cost=1000.0),
            AssignmentData(task_uid=2, resource_uid=1, cost=2000.0, actual_cost=1000.0),
        ],
    )

    comparison = compare_schedules(prior, later)
    cpm = compute_cpm(later)
    dcma = compute_dcma(later, cpm)
    manipulation = detect_manipulations(comparison, prior, later)
    ev = compute_earned_value(later)
    fa = analyze_float(comparison, prior, later)
    delay = analyze_delays(comparison, later)

    return {
        "prior_schedule": prior,
        "later_schedule": later,
        "comparison": comparison,
        "cpm": cpm,
        "dcma": dcma,
        "manipulation": manipulation,
        "earned_value": ev,
        "float_analysis": fa,
        "delay": delay,
    }


def _build_single_results() -> Dict:
    """Single-schedule pipeline (no comparison)."""
    tasks = [
        _make_task(
            1,
            "Design",
            duration=10.0,
            start=datetime(2026, 1, 5),
            finish=datetime(2026, 1, 16),
            baseline_start=datetime(2026, 1, 5),
            baseline_finish=datetime(2026, 1, 16),
            baseline_duration=10.0,
            percent_complete=100.0,
            total_slack=0.0,
            critical=True,
            successors=[2],
        ),
        _make_task(
            2,
            "Build",
            duration=20.0,
            start=datetime(2026, 1, 19),
            finish=datetime(2026, 2, 13),
            baseline_start=datetime(2026, 1, 19),
            baseline_finish=datetime(2026, 2, 13),
            baseline_duration=20.0,
            percent_complete=25.0,
            total_slack=0.0,
            critical=True,
            predecessors=[1],
        ),
    ]
    schedule = ScheduleData(
        project_info=ProjectInfo(
            name="Widget v1",
            status_date=datetime(2026, 1, 22),
            start_date=datetime(2026, 1, 5),
            finish_date=datetime(2026, 2, 13),
        ),
        tasks=tasks,
        relationships=[Relationship(predecessor_uid=1, successor_uid=2)],
    )
    cpm = compute_cpm(schedule)
    return {
        "prior_schedule": schedule,
        "later_schedule": None,
        "comparison": None,
        "cpm": cpm,
        "dcma": compute_dcma(schedule, cpm),
        "earned_value": compute_earned_value(schedule),
        "manipulation": None,
        "float_analysis": None,
        "delay": None,
    }


# --------------------------------------------------------------------------- #
# DOCX
# --------------------------------------------------------------------------- #


class TestDocxReport:
    def test_generates_valid_bytes(self):
        results = _build_comparison_results()
        data = generate_docx_report(results)
        assert isinstance(data, bytes)
        assert len(data) > 5000
        # .docx is a zip file, starts with PK\x03\x04
        assert data[:2] == b"PK"

    def test_round_trips_through_python_docx(self):
        data = generate_docx_report(_build_comparison_results())
        doc = Document(io.BytesIO(data))
        # Should have multiple headings
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Comparative Schedule Analysis" in text
        assert "Executive Summary" in text
        assert "DCMA 14-Point" in text
        assert "Manipulation Findings" in text
        assert "Critical Path Analysis" in text

    def test_ai_narrative_embedded(self):
        narrative = (
            "The project shows a five-day slip on the critical path.\n\n"
            "Weather is the dominant driver per the root-cause trace."
        )
        data = generate_docx_report(
            _build_comparison_results(),
            ai_narrative=narrative,
            analyst_name="J. Smith",
        )
        doc = Document(io.BytesIO(data))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "five-day slip" in text
        assert "Weather is the dominant driver" in text
        assert "J. Smith" in text

    def test_single_schedule_report(self):
        data = generate_docx_report(_build_single_results())
        doc = Document(io.BytesIO(data))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Schedule Analysis" in text
        # No comparison → should NOT include slippage section
        assert "Delay and Slippage Analysis" not in text

    def test_dcma_table_rows(self):
        data = generate_docx_report(_build_comparison_results())
        doc = Document(io.BytesIO(data))
        # Find the DCMA table — it has 5 columns and 15 rows (header + 14 metrics)
        found = False
        for table in doc.tables:
            if len(table.columns) == 5 and len(table.rows) == 15:
                hdr = [c.text for c in table.rows[0].cells]
                if hdr == ["#", "Metric", "Value", "Threshold", "Status"]:
                    found = True
                    break
        assert found, "DCMA table not found with expected shape"


# --------------------------------------------------------------------------- #
# XLSX
# --------------------------------------------------------------------------- #


class TestXlsxExport:
    def test_generates_valid_bytes(self):
        data = generate_xlsx_export(_build_comparison_results())
        assert isinstance(data, bytes)
        assert len(data) > 3000
        assert data[:2] == b"PK"

    def test_round_trips_through_openpyxl(self):
        data = generate_xlsx_export(_build_comparison_results())
        wb = load_workbook(io.BytesIO(data))
        expected_sheets = {
            "Summary",
            "All Tasks",
            "Critical Path",
            "Slippage",
            "DCMA Scorecard",
            "Manipulation",
        }
        assert expected_sheets.issubset(set(wb.sheetnames))

    def test_summary_has_key_rows(self):
        data = generate_xlsx_export(_build_comparison_results())
        wb = load_workbook(io.BytesIO(data))
        summary = wb["Summary"]
        values = {row[0].value: row[1].value for row in summary.iter_rows(min_row=2)}
        assert values.get("Project name") == "Bridge Replacement 42"
        # T3 slips 5 days, T4 cascades 5 days → tasks_slipped_count == 2
        assert values.get("Tasks slipped") == 2
        assert values.get("DCMA passed") is not None

    def test_all_tasks_has_every_task(self):
        data = generate_xlsx_export(_build_comparison_results())
        wb = load_workbook(io.BytesIO(data))
        tasks = wb["All Tasks"]
        # Header row + 4 tasks = 5 rows
        rows = list(tasks.iter_rows(values_only=True))
        assert len(rows) == 5
        uids = [row[0] for row in rows[1:]]
        assert set(uids) == {1, 2, 3, 4}

    def test_dcma_sheet_has_14_metrics(self):
        data = generate_xlsx_export(_build_comparison_results())
        wb = load_workbook(io.BytesIO(data))
        dcma = wb["DCMA Scorecard"]
        rows = list(dcma.iter_rows(values_only=True))
        # Header + 14 metrics
        assert len(rows) == 15

    def test_single_schedule_skips_comparison_sheets(self):
        data = generate_xlsx_export(_build_single_results())
        wb = load_workbook(io.BytesIO(data))
        assert "Slippage" not in wb.sheetnames
        assert "Manipulation" not in wb.sheetnames
        assert "All Tasks" in wb.sheetnames


# --------------------------------------------------------------------------- #
# PDF
# --------------------------------------------------------------------------- #


class TestPdfReport:
    def test_generates_valid_bytes(self):
        data = generate_pdf_report(_build_comparison_results())
        assert isinstance(data, bytes)
        assert len(data) > 3000
        assert data.startswith(b"%PDF"), "Missing PDF magic bytes"

    def test_pdf_has_trailer(self):
        data = generate_pdf_report(_build_comparison_results())
        # Every valid PDF ends with %%EOF
        assert b"%%EOF" in data[-64:]

    def test_ai_narrative_accepts_special_chars(self):
        # Make sure HTML-significant chars in the AI narrative don't break
        # reportlab's Paragraph parser.
        narrative = (
            "Project finish slipped <5 days>. Key driver: weather & soil.\n\n"
            "Second paragraph with ampersand & angle brackets <here>."
        )
        data = generate_pdf_report(
            _build_comparison_results(),
            ai_narrative=narrative,
            analyst_name="J. Smith",
        )
        assert data.startswith(b"%PDF")
        assert len(data) > 3000

    def test_single_schedule_pdf(self):
        data = generate_pdf_report(_build_single_results())
        assert data.startswith(b"%PDF")
        assert len(data) > 2000

    def test_empty_results_still_produces_pdf(self):
        # Degenerate case: mostly-empty results dict. The report should
        # still render without raising.
        data = generate_pdf_report({"comparison": None})
        assert data.startswith(b"%PDF")


# --------------------------------------------------------------------------- #
# Smoke: make sure the main.py route wiring still imports cleanly
# --------------------------------------------------------------------------- #


class TestMainRouteImports:
    def test_main_imports_exports(self):
        import app.main as main_mod

        assert hasattr(main_mod, "generate_docx_report")
        assert hasattr(main_mod, "generate_xlsx_export")
        assert hasattr(main_mod, "generate_pdf_report")
