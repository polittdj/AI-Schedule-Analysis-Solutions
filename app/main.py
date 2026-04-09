"""Flask application for the Schedule Forensics Local Tool.

Entry point: ``python -m app.main``

The app factory (`create_app`) wires together the routes and the
per-request config. Nothing here touches the JVM at import time — the
MPP parser starts the JVM lazily on the first real ``parse_mpp`` call
so ``python -m app.main`` boots quickly even without Java configured.

Storage model
-------------
Flask sessions are cookie-based (4 KB cap), far too small to hold a
full ``ScheduleData`` tree. Instead we store analysis artifacts on
disk in the configured ``UPLOAD_FOLDER`` as pickle files keyed by a
UUID, and the session only carries that UUID. A real deployment would
add a TTL reaper; Phase 4 keeps it simple.
"""
from __future__ import annotations

import io
import json
import pickle
import traceback
import uuid
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Optional

from flask import (
    Flask,
    Response,
    flash,
    jsonify,
    redirect,
    render_template,
    request,
    send_file,
    session,
    stream_with_context,
    url_for,
)
from werkzeug.utils import secure_filename

from app.ai.claude_client import ClaudeClient
from app.ai.ollama_client import OllamaClient
from app.ai.sanitizer import DataSanitizer
from app.config import Config, load_config
from app.engine.comparator import compare_schedules
from app.engine.cpm import compute_cpm
from app.engine.dcma import compute_dcma
from app.engine.delay_analysis import analyze_delays
from app.engine.earned_value import compute_earned_value
from app.engine.float_analysis import analyze_float
from app.engine.manipulation import detect_manipulations
from app.parser.schema import ScheduleData

# --------------------------------------------------------------------------- #
# Pipeline helpers
# --------------------------------------------------------------------------- #


def _run_single_schedule_pipeline(schedule: ScheduleData) -> Dict[str, Any]:
    """Run every single-schedule engine module on `schedule`."""
    cpm = compute_cpm(schedule)
    return {
        "cpm": cpm,
        "dcma": compute_dcma(schedule, cpm),
        "earned_value": compute_earned_value(schedule),
    }


def _run_comparative_pipeline(
    prior: ScheduleData, later: ScheduleData
) -> Dict[str, Any]:
    """Run the full comparative pipeline on `prior` + `later`."""
    later_single = _run_single_schedule_pipeline(later)
    comparison = compare_schedules(prior, later)
    return {
        **later_single,
        "comparison": comparison,
        "manipulation": detect_manipulations(comparison, prior, later),
        "float_analysis": analyze_float(comparison, prior, later),
        "delay": analyze_delays(comparison, later),
    }


def _assemble_results(
    prior: ScheduleData, later: Optional[ScheduleData]
) -> Dict[str, Any]:
    """Merge prior/later schedules with the computed engine output."""
    if later is None:
        base: Dict[str, Any] = {
            "prior_schedule": prior,
            "later_schedule": None,
            "comparison": None,
            "manipulation": None,
            "float_analysis": None,
            "delay": None,
            **_run_single_schedule_pipeline(prior),
        }
    else:
        base = {
            "prior_schedule": prior,
            "later_schedule": later,
            **_run_comparative_pipeline(prior, later),
        }
    base["generated_at"] = datetime.utcnow().isoformat()
    return base


# --------------------------------------------------------------------------- #
# Disk-backed analysis store
# --------------------------------------------------------------------------- #


def _store_path(upload_folder: Path, analysis_id: str) -> Path:
    return upload_folder / f"analysis_{analysis_id}.pkl"


def _save_analysis(upload_folder: Path, analysis_id: str, results: Dict[str, Any]) -> None:
    upload_folder.mkdir(parents=True, exist_ok=True)
    with _store_path(upload_folder, analysis_id).open("wb") as f:
        pickle.dump(results, f)


def _load_analysis(
    upload_folder: Path, analysis_id: str
) -> Optional[Dict[str, Any]]:
    path = _store_path(upload_folder, analysis_id)
    if not path.exists():
        return None
    with path.open("rb") as f:
        return pickle.load(f)


# --------------------------------------------------------------------------- #
# Serialization helpers for templates
# --------------------------------------------------------------------------- #


def _to_json_safe(obj: Any) -> Any:
    """Deep-convert Pydantic models and datetimes into JSON-friendly types."""
    if obj is None:
        return None
    if hasattr(obj, "model_dump"):
        return obj.model_dump(mode="json")
    if isinstance(obj, dict):
        return {k: _to_json_safe(v) for k, v in obj.items()}
    if isinstance(obj, list):
        return [_to_json_safe(v) for v in obj]
    if isinstance(obj, datetime):
        return obj.isoformat()
    return obj


def _results_for_template(results: Dict[str, Any]) -> Dict[str, Any]:
    return _to_json_safe(results)


# --------------------------------------------------------------------------- #
# File upload helpers
# --------------------------------------------------------------------------- #


def _valid_extension(filename: str, allowed: set[str]) -> bool:
    if "." not in filename:
        return False
    ext = filename.rsplit(".", 1)[1].lower()
    return ext in allowed


def _save_upload(file_storage, upload_folder: Path, prefix: str) -> Path:
    safe = secure_filename(file_storage.filename or "unnamed.mpp")
    unique = f"{prefix}_{uuid.uuid4().hex[:8]}_{safe}"
    upload_folder.mkdir(parents=True, exist_ok=True)
    target = upload_folder / unique
    file_storage.save(str(target))
    return target


# --------------------------------------------------------------------------- #
# Export helpers (DOCX / XLSX / PDF)
# --------------------------------------------------------------------------- #


def _primary_schedule(results: Dict[str, Any]) -> Optional[ScheduleData]:
    return results.get("later_schedule") or results.get("prior_schedule")


def _export_docx(results: Dict[str, Any]) -> io.BytesIO:
    from docx import Document

    doc = Document()
    doc.add_heading("Schedule Forensics Report", 0)
    doc.add_paragraph(f"Generated: {datetime.utcnow().isoformat()} UTC")

    schedule = _primary_schedule(results)
    if schedule is not None:
        info = schedule.project_info
        doc.add_heading("Project Overview", 1)
        doc.add_paragraph(f"Name: {info.name or '—'}")
        doc.add_paragraph(f"Status date: {info.status_date or '—'}")
        doc.add_paragraph(f"Start / Finish: {info.start_date or '—'} → {info.finish_date or '—'}")
        doc.add_paragraph(f"Task count: {len(schedule.tasks)}")

    dcma = results.get("dcma")
    if dcma is not None:
        doc.add_heading("DCMA 14-Point Assessment", 1)
        doc.add_paragraph(
            f"Overall: {dcma.passed_count} passed / {dcma.failed_count} failed "
            f"({dcma.overall_score_pct:.1f}%)."
        )
        table = doc.add_table(rows=1, cols=5)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "#"
        hdr[1].text = "Metric"
        hdr[2].text = "Value"
        hdr[3].text = "Threshold"
        hdr[4].text = "Status"
        for m in dcma.metrics:
            row = table.add_row().cells
            row[0].text = str(m.number)
            row[1].text = m.name
            row[2].text = f"{m.value}{m.unit}" if m.unit == "%" else f"{m.value}"
            row[3].text = f"{m.comparison}{m.threshold}"
            row[4].text = "PASS" if m.passed else "FAIL"

    manipulation = results.get("manipulation")
    if manipulation is not None:
        doc.add_heading("Manipulation Findings", 1)
        doc.add_paragraph(
            f"Overall score: {manipulation.overall_score:.1f}/100 "
            f"(HIGH={manipulation.confidence_summary.get('HIGH', 0)}, "
            f"MEDIUM={manipulation.confidence_summary.get('MEDIUM', 0)}, "
            f"LOW={manipulation.confidence_summary.get('LOW', 0)})"
        )
        for f in manipulation.findings[:50]:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"[{f.confidence}] [{f.category}] ").bold = True
            p.add_run(f"{f.description}")

    comparison = results.get("comparison")
    if comparison is not None and comparison.task_deltas:
        doc.add_heading("Top 10 Slipped Tasks", 1)
        top = sorted(
            comparison.task_deltas,
            key=lambda d: d.finish_slip_days or 0.0,
            reverse=True,
        )[:10]
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "UID"
        hdr[1].text = "Task"
        hdr[2].text = "Finish Slip (d)"
        hdr[3].text = "Duration Δ (d)"
        for d in top:
            row = table.add_row().cells
            row[0].text = str(d.uid)
            row[1].text = d.name or "—"
            row[2].text = f"{d.finish_slip_days or 0:.1f}"
            row[3].text = f"{d.duration_change_days or 0:.1f}"

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _export_xlsx(results: Dict[str, Any]) -> io.BytesIO:
    from openpyxl import Workbook

    wb = Workbook()
    ws_tasks = wb.active
    ws_tasks.title = "Tasks"
    ws_tasks.append(
        [
            "UID",
            "ID",
            "Name",
            "WBS",
            "Start",
            "Finish",
            "Duration (d)",
            "% Complete",
            "Total Slack (d)",
            "Critical",
            "Summary",
            "Milestone",
        ]
    )
    schedule = _primary_schedule(results)
    if schedule is not None:
        for t in schedule.tasks:
            ws_tasks.append(
                [
                    t.uid,
                    t.id,
                    t.name,
                    t.wbs,
                    t.start.isoformat() if t.start else None,
                    t.finish.isoformat() if t.finish else None,
                    t.duration,
                    t.percent_complete,
                    t.total_slack,
                    "Y" if t.critical else "",
                    "Y" if t.summary else "",
                    "Y" if t.milestone else "",
                ]
            )

    dcma = results.get("dcma")
    if dcma is not None:
        ws_dcma = wb.create_sheet("DCMA")
        ws_dcma.append(["#", "Metric", "Value", "Threshold", "Status"])
        for m in dcma.metrics:
            ws_dcma.append(
                [
                    m.number,
                    m.name,
                    m.value,
                    f"{m.comparison}{m.threshold}",
                    "PASS" if m.passed else "FAIL",
                ]
            )

    comparison = results.get("comparison")
    if comparison is not None:
        ws_deltas = wb.create_sheet("Task Deltas")
        ws_deltas.append(
            [
                "UID",
                "Name",
                "Start Slip (d)",
                "Finish Slip (d)",
                "Duration Δ (d)",
                "Total Slack Δ (d)",
            ]
        )
        for d in comparison.task_deltas:
            ws_deltas.append(
                [
                    d.uid,
                    d.name,
                    d.start_slip_days,
                    d.finish_slip_days,
                    d.duration_change_days,
                    d.total_slack_delta,
                ]
            )

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf


def _export_pdf(results: Dict[str, Any]) -> io.BytesIO:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import (
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter)
    styles = getSampleStyleSheet()
    story = [Paragraph("Schedule Forensics Report", styles["Title"])]
    story.append(Spacer(1, 12))
    story.append(
        Paragraph(
            f"Generated: {datetime.utcnow().isoformat()} UTC", styles["Normal"]
        )
    )
    story.append(Spacer(1, 12))

    schedule = _primary_schedule(results)
    if schedule is not None:
        info = schedule.project_info
        story.append(Paragraph("Project Overview", styles["Heading2"]))
        story.append(Paragraph(f"Name: {info.name or '—'}", styles["Normal"]))
        story.append(
            Paragraph(f"Status date: {info.status_date or '—'}", styles["Normal"])
        )
        story.append(
            Paragraph(
                f"Tasks: {len(schedule.tasks)}", styles["Normal"]
            )
        )
        story.append(Spacer(1, 12))

    dcma = results.get("dcma")
    if dcma is not None:
        story.append(Paragraph("DCMA 14-Point Assessment", styles["Heading2"]))
        data = [["#", "Metric", "Value", "Threshold", "Status"]]
        for m in dcma.metrics:
            data.append(
                [
                    str(m.number),
                    m.name,
                    f"{m.value}{m.unit}" if m.unit == "%" else f"{m.value}",
                    f"{m.comparison}{m.threshold}",
                    "PASS" if m.passed else "FAIL",
                ]
            )
        t = Table(data, repeatRows=1)
        t.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1b2432")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                ]
            )
        )
        story.append(t)

    doc.build(story)
    buf.seek(0)
    return buf


# --------------------------------------------------------------------------- #
# App factory
# --------------------------------------------------------------------------- #


def create_app(config: Optional[Config] = None) -> Flask:
    cfg = config or load_config()

    app = Flask(
        __name__,
        template_folder="web/templates",
        static_folder="web/static",
        static_url_path="/static",
    )
    app.config["SECRET_KEY"] = cfg.SECRET_KEY
    app.config["MAX_CONTENT_LENGTH"] = cfg.MAX_FILE_SIZE
    app.config["APP_CONFIG"] = cfg

    cfg.UPLOAD_FOLDER.mkdir(parents=True, exist_ok=True)

    # ------------------------------------------------------------------ #
    # Routes
    # ------------------------------------------------------------------ #

    @app.context_processor
    def inject_config() -> Dict[str, Any]:
        return {"config": cfg}

    @app.route("/")
    def index():
        return render_template("index.html")

    @app.route("/analyze", methods=["POST"])
    def analyze():
        mode = request.form.get("mode", "single")
        prior_file = request.files.get("prior_file")
        later_file = request.files.get("later_file")

        if prior_file is None or not (prior_file.filename or "").strip():
            flash("Please upload a prior schedule file.", "error")
            return redirect(url_for("index"))
        if not _valid_extension(prior_file.filename, set(cfg.ALLOWED_EXTENSIONS)):
            flash(
                f"Unsupported file type. Allowed: {', '.join(sorted(cfg.ALLOWED_EXTENSIONS))}.",
                "error",
            )
            return redirect(url_for("index"))

        if mode == "comparative":
            if later_file is None or not (later_file.filename or "").strip():
                flash("Comparative mode requires two files.", "error")
                return redirect(url_for("index"))
            if not _valid_extension(later_file.filename, set(cfg.ALLOWED_EXTENSIONS)):
                flash("Unsupported later-file type.", "error")
                return redirect(url_for("index"))

        try:
            from app.parser.mpp_reader import parse_mpp  # lazy: avoids JVM at import
        except Exception as exc:
            flash(f"Parser unavailable: {exc}", "error")
            return redirect(url_for("index"))

        try:
            prior_path = _save_upload(prior_file, cfg.UPLOAD_FOLDER, "prior")
            prior_schedule = parse_mpp(str(prior_path))

            later_schedule: Optional[ScheduleData] = None
            if mode == "comparative" and later_file is not None:
                later_path = _save_upload(later_file, cfg.UPLOAD_FOLDER, "later")
                later_schedule = parse_mpp(str(later_path))

            results = _assemble_results(prior_schedule, later_schedule)
        except Exception as exc:
            app.logger.error("Analysis failed: %s", exc)
            app.logger.error(traceback.format_exc())
            flash(f"Analysis failed: {exc}", "error")
            return redirect(url_for("index"))

        analysis_id = uuid.uuid4().hex
        _save_analysis(cfg.UPLOAD_FOLDER, analysis_id, results)
        session["analysis_id"] = analysis_id
        flash("Analysis complete.", "success")
        return redirect(url_for("analysis_page"))

    @app.route("/analysis")
    def analysis_page():
        analysis_id = session.get("analysis_id")
        if not analysis_id:
            flash("No analysis in progress. Upload a schedule to get started.", "info")
            return redirect(url_for("index"))
        results = _load_analysis(cfg.UPLOAD_FOLDER, analysis_id)
        if results is None:
            session.pop("analysis_id", None)
            flash("Analysis artifacts were purged. Please re-upload.", "info")
            return redirect(url_for("index"))

        template_ctx = _results_for_template(results)
        return render_template(
            "analysis.html",
            results=template_ctx,
            results_json=json.dumps(template_ctx, default=str),
            has_comparison=results.get("comparison") is not None,
        )

    @app.route("/ai-analyze", methods=["POST"])
    def ai_analyze():
        analysis_id = session.get("analysis_id")
        if not analysis_id:
            return jsonify({"error": "No analysis loaded"}), 400
        results = _load_analysis(cfg.UPLOAD_FOLDER, analysis_id)
        if results is None:
            return jsonify({"error": "Analysis artifacts purged"}), 404

        user_request = request.form.get(
            "request", "Provide a comprehensive forensic analysis."
        )
        backend_name = request.form.get("backend") or cfg.AI_MODE
        if backend_name == "cloud":
            backend = ClaudeClient(
                api_key=cfg.ANTHROPIC_API_KEY, model=cfg.ANTHROPIC_MODEL
            )
        else:
            backend = OllamaClient(
                url=cfg.OLLAMA_URL,
                model=cfg.OLLAMA_MODEL,
                timeout=cfg.OLLAMA_TIMEOUT,
            )

        if not backend.is_available():
            return (
                jsonify(
                    {
                        "error": f"AI backend '{backend_name}' is not available. "
                        "Check settings."
                    }
                ),
                503,
            )

        sanitizer = DataSanitizer() if cfg.SANITIZE_DATA else None
        engine_results = results
        if sanitizer is not None:
            engine_results = sanitizer.sanitize(results)

        def generate():
            try:
                for chunk in backend.stream_analyze(engine_results, user_request):
                    if sanitizer is not None and chunk:
                        chunk = sanitizer.desanitize_text(chunk)
                    if chunk:
                        yield f"data: {json.dumps({'chunk': chunk})}\n\n"
                yield f"data: {json.dumps({'done': True})}\n\n"
            except Exception as exc:
                yield f"data: {json.dumps({'error': str(exc)})}\n\n"

        return Response(
            stream_with_context(generate()),
            mimetype="text/event-stream",
            headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
        )

    @app.route("/export/<string:fmt>")
    def export(fmt: str):
        analysis_id = session.get("analysis_id")
        if not analysis_id:
            flash("No analysis to export.", "error")
            return redirect(url_for("index"))
        results = _load_analysis(cfg.UPLOAD_FOLDER, analysis_id)
        if results is None:
            flash("Analysis artifacts were purged. Please re-upload.", "error")
            return redirect(url_for("index"))

        try:
            if fmt == "docx":
                buf = _export_docx(results)
                return send_file(
                    buf,
                    mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    as_attachment=True,
                    download_name="schedule_forensics_report.docx",
                )
            if fmt == "xlsx":
                buf = _export_xlsx(results)
                return send_file(
                    buf,
                    mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    as_attachment=True,
                    download_name="schedule_forensics_data.xlsx",
                )
            if fmt == "pdf":
                buf = _export_pdf(results)
                return send_file(
                    buf,
                    mimetype="application/pdf",
                    as_attachment=True,
                    download_name="schedule_forensics_report.pdf",
                )
        except Exception as exc:
            app.logger.error("Export failed: %s", exc)
            flash(f"Export failed: {exc}", "error")
            return redirect(url_for("analysis_page"))

        flash(f"Unsupported export format: {fmt}", "error")
        return redirect(url_for("analysis_page"))

    @app.route("/settings", methods=["GET", "POST"])
    def settings_page():
        # Probe Ollama status for the UI indicator
        ollama_up = OllamaClient(url=cfg.OLLAMA_URL, model=cfg.OLLAMA_MODEL).is_available()
        claude_up = ClaudeClient(api_key=cfg.ANTHROPIC_API_KEY).is_available()

        if request.method == "POST":
            # Settings changes are process-ephemeral in Phase 4. A real
            # deployment would persist to a local .env / secret store.
            flash(
                "Phase 4 note: settings are read from environment variables and "
                "cannot be persisted from the UI yet. Set env vars and restart.",
                "info",
            )
            return redirect(url_for("settings_page"))

        return render_template(
            "settings.html",
            ollama_up=ollama_up,
            claude_up=claude_up,
            cfg_snapshot=cfg.as_dict(),
        )

    # ------------------------------------------------------------------ #
    # Error handlers
    # ------------------------------------------------------------------ #

    @app.errorhandler(413)
    def too_large(_exc):
        flash(
            f"File exceeds the {cfg.MAX_FILE_SIZE // (1024 * 1024)} MB upload limit.",
            "error",
        )
        return redirect(url_for("index")), 302

    @app.errorhandler(404)
    def not_found(_exc):
        return render_template("base.html"), 404

    @app.errorhandler(500)
    def server_error(exc):
        app.logger.error("Server error: %s", exc)
        flash("Unexpected server error. Check the console log.", "error")
        return redirect(url_for("index")), 302

    return app


app = create_app()


if __name__ == "__main__":
    app.run(host="127.0.0.1", port=5000, debug=False)
