"""Word document report generator.

Produces a professional forensic-analysis report in .docx format from
the engine results dict plus an optional AI narrative. The report is
self-contained — the same file is suitable for emailing to owners,
attaching to a claim package, or printing to PDF through Word.

Design notes
------------
* Everything uses Arial 11 pt for body text and Arial for headings to
  match the visual weight of most claim reports.
* DCMA pass/fail cells are color-shaded by direct XML manipulation
  because python-docx doesn't expose a cell-background API.
* Manipulation findings are color-coded by confidence on the
  confidence tag itself, not the whole paragraph, so the body text
  stays readable in black.
* Returns ``bytes`` (not ``BytesIO``) so callers can wrap it in
  whatever stream they need. The Flask route wraps it back in a
  ``BytesIO`` for ``send_file``.
"""
from __future__ import annotations

import io
from datetime import datetime
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# Palette matches the web UI dark theme for visual consistency.
HDR_FILL = "1B2432"        # table header background
PASS_FILL = "D4F5E3"       # pastel green
FAIL_FILL = "FBD7DB"       # pastel red
HIGH_COLOR = RGBColor(0xFF, 0x5C, 0x6C)
MEDIUM_COLOR = RGBColor(0xF5, 0xB8, 0x3D)
LOW_COLOR = RGBColor(0x4A, 0x9E, 0xFF)

BODY_FONT = "Arial"


# --------------------------------------------------------------------------- #
# Public entry point
# --------------------------------------------------------------------------- #


def generate_docx_report(
    results: Dict[str, Any],
    ai_narrative: Optional[str] = None,
    analyst_name: Optional[str] = None,
) -> bytes:
    """Render the full Word report and return it as bytes."""
    doc = Document()
    _apply_body_font(doc)

    has_comparison = results.get("comparison") is not None
    schedule = _primary_schedule(results)

    title = "Comparative Schedule Analysis" if has_comparison else "Schedule Analysis"
    _add_title_page(doc, title, schedule, analyst_name)
    _add_header_footer(doc, title)

    doc.add_page_break()
    _add_executive_summary(doc, ai_narrative)

    if results.get("dcma") is not None:
        _add_dcma_section(doc, results["dcma"])

    if results.get("cpm") is not None:
        _add_critical_path_section(doc, results["cpm"], schedule)

    if has_comparison:
        _add_slippage_section(
            doc, results["comparison"], results.get("delay")
        )

    if results.get("manipulation") is not None:
        _add_manipulation_section(doc, results["manipulation"])

    if results.get("float_analysis") is not None:
        _add_float_section(doc, results["float_analysis"])

    if results.get("earned_value") is not None:
        _add_earned_value_section(doc, results["earned_value"])

    if results.get("trend") is not None:
        _add_trend_section(doc, results["trend"])

    _add_recommendations(doc, results)
    _add_conclusion(doc, results, has_comparison)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# --------------------------------------------------------------------------- #
# Styling helpers (direct XML where python-docx has no API)
# --------------------------------------------------------------------------- #


def _apply_body_font(doc: Document) -> None:
    """Force Arial on Normal + heading styles."""
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(11)
    # Also set the East Asian / Complex Script font so Word doesn't
    # silently swap us back to Calibri on non-ASCII text.
    rPr = normal.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), BODY_FONT)

    for level in (1, 2, 3):
        try:
            hstyle = doc.styles[f"Heading {level}"]
            hstyle.font.name = BODY_FONT
        except KeyError:
            continue


def _set_cell_shading(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _style_header_row(row, headers: List[str]) -> None:
    for i, text in enumerate(headers):
        cell = row.cells[i]
        cell.text = text
        _set_cell_shading(cell, HDR_FILL)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10)


def _add_header_footer(doc: Document, title: str) -> None:
    section = doc.sections[0]
    header_para = section.header.paragraphs[0]
    header_para.text = title
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header_para.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x5A, 0x65, 0x78)

    footer_para = section.footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.text = ""
    run = footer_para.add_run("Page ")
    run.font.size = Pt(9)
    _append_page_field(footer_para)
    footer_para.add_run(" of ").font.size = Pt(9)
    _append_page_count_field(footer_para)


def _append_page_field(paragraph) -> None:
    run = paragraph.add_run()
    run.font.size = Pt(9)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def _append_page_count_field(paragraph) -> None:
    run = paragraph.add_run()
    run.font.size = Pt(9)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "NUMPAGES"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


# --------------------------------------------------------------------------- #
# Utilities
# --------------------------------------------------------------------------- #


def _primary_schedule(results: Dict[str, Any]):
    return results.get("later_schedule") or results.get("prior_schedule")


def _fmt_date(value: Any) -> str:
    if value is None:
        return "—"
    if hasattr(value, "strftime"):
        try:
            return value.strftime("%Y-%m-%d")
        except Exception:
            pass
    return str(value)


def _fmt_signed(value: Optional[float], suffix: str = "") -> str:
    if value is None:
        return "—"
    return f"{value:+.1f}{suffix}"


def _fmt_num(value: Optional[float], decimals: int = 1, suffix: str = "") -> str:
    if value is None:
        return "—"
    return f"{value:.{decimals}f}{suffix}"


# --------------------------------------------------------------------------- #
# Sections
# --------------------------------------------------------------------------- #


def _add_title_page(
    doc: Document, title: str, schedule, analyst_name: Optional[str]
) -> None:
    doc.add_heading(title, 0)

    if schedule is not None:
        info = schedule.project_info
        p = doc.add_paragraph()
        run = p.add_run(info.name or "Unnamed Project")
        run.font.size = Pt(16)
        run.font.bold = True

        meta = doc.add_paragraph()
        meta.add_run("Start: ").font.bold = True
        meta.add_run(f"{_fmt_date(info.start_date)}\n")
        meta.add_run("Finish: ").font.bold = True
        meta.add_run(f"{_fmt_date(info.finish_date)}\n")
        meta.add_run("Status Date: ").font.bold = True
        meta.add_run(f"{_fmt_date(info.status_date)}")

    if analyst_name:
        p = doc.add_paragraph()
        p.add_run("Analyst: ").font.bold = True
        p.add_run(analyst_name)

    gen = doc.add_paragraph()
    gen.add_run("Generated: ").font.bold = True
    gen.add_run(f"{datetime.utcnow().isoformat(timespec='seconds')} UTC")


def _add_executive_summary(doc: Document, ai_narrative: Optional[str]) -> None:
    doc.add_heading("Executive Summary", 1)
    if ai_narrative and ai_narrative.strip():
        for para in ai_narrative.strip().split("\n\n"):
            text = para.strip()
            if text:
                doc.add_paragraph(text)
    else:
        p = doc.add_paragraph()
        p.add_run(
            "AI narrative not yet generated. Open the Analysis tab in the web "
            "UI and click 'Generate AI Analysis' to produce a narrative, then "
            "re-export this report to include it."
        ).italic = True


def _add_dcma_section(doc: Document, dcma) -> None:
    doc.add_heading("Schedule Health Assessment — DCMA 14-Point", 1)

    summary = doc.add_paragraph()
    summary.add_run(
        f"Overall: {dcma.passed_count} passed / {dcma.failed_count} failed "
        f"({dcma.overall_score_pct:.1f}%)"
    ).bold = True

    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    _style_header_row(table.rows[0], ["#", "Metric", "Value", "Threshold", "Status"])

    for m in dcma.metrics:
        row = table.add_row()
        row.cells[0].text = str(m.number)
        row.cells[1].text = m.name
        if m.unit == "%":
            row.cells[2].text = f"{m.value:.2f}%"
        elif m.unit == "index":
            row.cells[2].text = f"{m.value:.3f}"
        else:
            row.cells[2].text = str(m.value)
        row.cells[3].text = f"{m.comparison}{m.threshold}"
        status_cell = row.cells[4]
        status_cell.text = "PASS" if m.passed else "FAIL"
        _set_cell_shading(status_cell, PASS_FILL if m.passed else FAIL_FILL)
        for para in status_cell.paragraphs:
            for run in para.runs:
                run.font.bold = True


def _add_critical_path_section(doc: Document, cpm, schedule) -> None:
    doc.add_heading("Critical Path Analysis", 1)
    summary = doc.add_paragraph()
    summary.add_run(
        f"Project duration: {cpm.project_duration_days:.1f} working days — "
        f"{len(cpm.critical_path_uids)} task(s) on the critical path."
    )

    if not cpm.critical_path_uids or schedule is None:
        return

    task_by_uid = {t.uid: t for t in schedule.tasks}
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    _style_header_row(
        table.rows[0], ["Seq", "Task", "Finish", "Dur (d)", "Total Float (d)"]
    )
    float_lookup = cpm.task_floats or {}

    for idx, uid in enumerate(cpm.critical_path_uids[:30], start=1):
        task = task_by_uid.get(uid)
        if task is None:
            continue
        tf = float_lookup.get(uid)
        total_float = getattr(tf, "total_float", None) if tf else None
        row = table.add_row()
        row.cells[0].text = str(idx)
        row.cells[1].text = task.name or f"Task {uid}"
        row.cells[2].text = _fmt_date(task.finish)
        row.cells[3].text = _fmt_num(task.duration)
        row.cells[4].text = _fmt_num(total_float)


def _add_slippage_section(doc: Document, comparison, delay) -> None:
    doc.add_heading("Delay and Slippage Analysis", 1)

    stats = doc.add_paragraph()
    stats.add_run(
        f"Tasks slipped: {comparison.tasks_slipped_count}   "
        f"Pulled in: {comparison.tasks_pulled_in_count}   "
        f"Completed: {comparison.tasks_completed_count}   "
        f"Added: {comparison.tasks_added_count}   "
        f"Deleted: {comparison.tasks_deleted_count}"
    ).bold = True
    if comparison.completion_date_slip_days is not None:
        doc.add_paragraph(
            f"Project completion slip: "
            f"{comparison.completion_date_slip_days:+.1f} calendar days."
        )

    # Top slips
    slipped = [
        d
        for d in comparison.task_deltas
        if d.finish_slip_days and d.finish_slip_days > 0
    ]
    slipped.sort(key=lambda d: d.finish_slip_days or 0, reverse=True)

    if slipped:
        doc.add_heading("Top 10 Slipped Tasks", 2)
        table = doc.add_table(rows=1, cols=5)
        table.style = "Light Grid Accent 1"
        _style_header_row(
            table.rows[0],
            ["UID", "Task", "Start Slip (d)", "Finish Slip (d)", "Δ Dur (d)"],
        )
        for d in slipped[:10]:
            row = table.add_row()
            row.cells[0].text = str(d.uid)
            row.cells[1].text = d.name or "—"
            row.cells[2].text = _fmt_signed(d.start_slip_days)
            row.cells[3].text = _fmt_signed(d.finish_slip_days)
            row.cells[4].text = _fmt_signed(d.duration_change_days)

    # Duration changes
    duration_changes = [
        d
        for d in comparison.task_deltas
        if d.duration_change_days and d.duration_change_days != 0
    ]
    if duration_changes:
        doc.add_heading("Duration Changes", 2)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        _style_header_row(table.rows[0], ["UID", "Task", "Δ Duration (d)"])
        top_dur = sorted(
            duration_changes,
            key=lambda d: abs(d.duration_change_days or 0),
            reverse=True,
        )[:10]
        for d in top_dur:
            row = table.add_row()
            row.cells[0].text = str(d.uid)
            row.cells[1].text = d.name or "—"
            row.cells[2].text = _fmt_signed(d.duration_change_days)

    # Root cause narrative
    if delay is not None:
        doc.add_heading("Root Cause Analysis", 2)
        if delay.first_mover_uid is not None:
            p = doc.add_paragraph()
            p.add_run("First mover: ").bold = True
            name = delay.first_mover_name or f"Task {delay.first_mover_uid}"
            p.add_run(
                f"{name} — slipped "
                f"{delay.first_mover_slip_days:.1f} days."
            )
        if delay.root_causes:
            doc.add_paragraph("Top delay drivers:")
            for cause in delay.root_causes[:10]:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(f"[{cause.category}] ").bold = True
                task_label = cause.task_name or f"Task {cause.task_uid}"
                p.add_run(f"{task_label} — {cause.slip_days:.1f} days")
                if cause.on_critical_path:
                    p.add_run(" (critical path)")


def _add_manipulation_section(doc: Document, manip) -> None:
    doc.add_heading("Manipulation Findings", 1)

    summary = doc.add_paragraph()
    summary.add_run(f"Overall score: {manip.overall_score:.0f}/100").bold = True
    summary.add_run(
        f"   (HIGH: {manip.confidence_summary.get('HIGH', 0)}, "
        f"MEDIUM: {manip.confidence_summary.get('MEDIUM', 0)}, "
        f"LOW: {manip.confidence_summary.get('LOW', 0)})"
    )

    if not manip.findings:
        doc.add_paragraph("No manipulation patterns detected.")
        return

    conf_rank = {"HIGH": 0, "MEDIUM": 1, "LOW": 2}
    ordered = sorted(
        manip.findings,
        key=lambda f: (conf_rank.get(f.confidence, 9), -f.severity_score),
    )
    for finding in ordered:
        p = doc.add_paragraph(style="List Bullet")
        conf_run = p.add_run(f"[{finding.confidence}] ")
        conf_run.bold = True
        if finding.confidence == "HIGH":
            conf_run.font.color.rgb = HIGH_COLOR
        elif finding.confidence == "MEDIUM":
            conf_run.font.color.rgb = MEDIUM_COLOR
        else:
            conf_run.font.color.rgb = LOW_COLOR
        cat_run = p.add_run(f"[{finding.category}] ")
        cat_run.bold = True
        p.add_run(finding.description)
        if finding.task_name:
            p.add_run(f" — Task: {finding.task_name}")


def _add_float_section(doc: Document, fa) -> None:
    doc.add_heading("Float Analysis", 1)
    p = doc.add_paragraph()
    p.add_run("Trend: ").bold = True
    p.add_run(fa.trend)
    doc.add_paragraph(
        f"Net float delta: {fa.net_float_delta:+.1f} days across "
        f"{len(fa.task_changes)} compared tasks."
    )
    doc.add_paragraph(
        f"Tasks that became critical: {len(fa.became_critical_uids)}"
    )
    doc.add_paragraph(
        f"Tasks that dropped off the critical path: "
        f"{len(fa.dropped_off_critical_uids)}"
    )


def _add_earned_value_section(doc: Document, ev) -> None:
    doc.add_heading("Earned Value Summary", 1)
    doc.add_paragraph(f"Units: {ev.units}")

    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    _style_header_row(table.rows[0], ["Metric", "Value"])

    rows = [
        ("BAC — Budget at Completion", _fmt_num(ev.budget_at_completion)),
        ("PV — Planned Value", _fmt_num(ev.planned_value)),
        ("EV — Earned Value", _fmt_num(ev.earned_value)),
        ("SV — Schedule Variance", _fmt_signed(ev.schedule_variance)),
        (
            "SPI — Schedule Performance Index",
            _fmt_num(ev.schedule_performance_index, decimals=3),
        ),
    ]
    if ev.actual_cost is not None:
        rows.append(("AC — Actual Cost", _fmt_num(ev.actual_cost)))
    if ev.cost_variance is not None:
        rows.append(("CV — Cost Variance", _fmt_signed(ev.cost_variance)))
    if ev.cost_performance_index is not None:
        rows.append(
            (
                "CPI — Cost Performance Index",
                _fmt_num(ev.cost_performance_index, decimals=3),
            )
        )
    if ev.estimate_at_completion is not None:
        rows.append(
            ("EAC — Estimate at Completion", _fmt_num(ev.estimate_at_completion))
        )

    for label, value in rows:
        r = table.add_row()
        r.cells[0].text = label
        r.cells[1].text = value


def _add_trend_section(doc: Document, trend) -> None:
    doc.add_heading("Trend Analysis (multi-update time-series)", 1)

    summary = doc.add_paragraph()
    summary.add_run(
        f"Updates analyzed: {trend.update_count}   "
        f"Float: {trend.float_trend}   "
        f"SPI: {trend.spi_trend}   "
        f"Manipulation: {trend.manipulation_trend}"
    ).bold = True
    if trend.completion_date_drift_days is not None:
        doc.add_paragraph(
            f"Completion drift (first → last): "
            f"{trend.completion_date_drift_days:+.1f} calendar days."
        )
    if trend.narrative:
        doc.add_paragraph(trend.narrative)

    if trend.data_points:
        doc.add_heading("Per-Update Metrics", 2)
        table = doc.add_table(rows=1, cols=6)
        table.style = "Light Grid Accent 1"
        _style_header_row(
            table.rows[0],
            ["Update", "Status Date", "Finish", "SPI", "Manip", "Slip (d)"],
        )
        for dp in trend.data_points:
            row = table.add_row()
            row.cells[0].text = dp.update_label
            row.cells[1].text = _fmt_date(dp.status_date)
            row.cells[2].text = _fmt_date(dp.project_finish)
            row.cells[3].text = _fmt_num(dp.spi, decimals=3)
            row.cells[4].text = _fmt_num(dp.manipulation_score)
            row.cells[5].text = _fmt_signed(dp.finish_slip_since_prior_days)

    if trend.task_compressions:
        doc.add_heading("Cumulative Task Compressions", 2)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        _style_header_row(
            table.rows[0], ["UID", "Task", "Cumulative Δ (d)", "Events"]
        )
        for tc in trend.task_compressions:
            row = table.add_row()
            row.cells[0].text = str(tc.uid)
            row.cells[1].text = tc.name or "—"
            row.cells[2].text = _fmt_signed(tc.cumulative_duration_change_days)
            row.cells[3].text = str(tc.compression_events)

    if trend.baseline_resets:
        doc.add_heading("Baseline Reset Events", 2)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        _style_header_row(
            table.rows[0], ["Update", "Affected Tasks", "Max Shift (d)"]
        )
        for ev in trend.baseline_resets:
            row = table.add_row()
            row.cells[0].text = ev.update_label
            row.cells[1].text = str(ev.affected_task_count)
            row.cells[2].text = _fmt_num(ev.max_baseline_shift_days)


def _add_recommendations(doc: Document, results: Dict[str, Any]) -> None:
    doc.add_heading("Opportunities and Recommendations", 1)
    for rec in _build_recommendations(results):
        doc.add_paragraph(rec, style="List Bullet")


def _add_conclusion(
    doc: Document, results: Dict[str, Any], has_comparison: bool
) -> None:
    doc.add_heading("Conclusion", 1)
    doc.add_paragraph(_build_conclusion_text(results, has_comparison))


def _build_recommendations(results: Dict[str, Any]) -> List[str]:
    recs: List[str] = []
    dcma = results.get("dcma")
    if dcma is not None:
        failed = [m for m in dcma.metrics if not m.passed]
        for m in failed[:5]:
            recs.append(
                f"DCMA #{m.number} ({m.name}) is failing at {m.value}{m.unit}. "
                f"Target: {m.comparison}{m.threshold}. Remediate before next update."
            )
    manip = results.get("manipulation")
    if manip is not None and manip.overall_score >= 20:
        recs.append(
            f"Manipulation score of {manip.overall_score:.0f}/100 warrants "
            "independent review of the schedule update log."
        )
    fa = results.get("float_analysis")
    if fa is not None and fa.trend == "consuming":
        recs.append(
            f"Float consumption trend is negative "
            f"({fa.net_float_delta:+.1f} d net). Review mitigation options for "
            f"the {len(fa.became_critical_uids)} task(s) that became critical."
        )
    ev = results.get("earned_value")
    if ev is not None and ev.schedule_performance_index < 0.95:
        recs.append(
            f"SPI of {ev.schedule_performance_index:.2f} indicates schedule "
            "performance below plan. Recommend recovery-plan review."
        )
    if not recs:
        recs.append(
            "No critical issues detected. Continue weekly schedule updates "
            "and re-run DCMA checks after each update."
        )
    return recs


def _build_conclusion_text(
    results: Dict[str, Any], has_comparison: bool
) -> str:
    parts: List[str] = []
    dcma = results.get("dcma")
    if dcma is not None:
        parts.append(
            f"The schedule passes {dcma.passed_count} of "
            f"{len(dcma.metrics)} DCMA metrics "
            f"({dcma.overall_score_pct:.0f}% overall health)."
        )
    if has_comparison:
        comparison = results.get("comparison")
        if comparison is not None:
            parts.append(
                f"Between the prior and later updates, "
                f"{comparison.tasks_slipped_count} task(s) slipped, "
                f"{comparison.tasks_completed_count} completed, and "
                f"{comparison.tasks_added_count} new task(s) were added."
            )
        manip = results.get("manipulation")
        if manip is not None:
            parts.append(
                f"Manipulation indicators scored "
                f"{manip.overall_score:.0f}/100 "
                f"with {manip.confidence_summary.get('HIGH', 0)} HIGH-confidence "
                "finding(s)."
            )
    else:
        parts.append(
            "This is a single-schedule snapshot; run a comparative analysis "
            "against a prior update to assess slippage and manipulation."
        )
    if not parts:
        parts.append("Report generated with no engine findings available.")
    return " ".join(parts)
